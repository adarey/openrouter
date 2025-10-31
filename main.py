# ========================================
# 📦 IMPORTS TEST
# ========================================
import streamlit as st
import requests
from io import BytesIO
from datetime import datetime
from pathlib import Path
from docx import Document
from fpdf import FPDF
# ========================================

st.set_page_config(page_title="OpenRouter • IA gratuite", page_icon="🧩")
st.title("😋 Modèles gratuits – OpenRouter")
st.markdown("Liste automatisée depuis l’API officielle OpenRouter.\nSélectionne, envoie un prompt, reçois la réponse !")

# -------------------------------
# 🔧 Dossier d’exports automatiques
# -------------------------------
EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)

# -------------------------------
# 🔧 Fonctions d’export
# -------------------------------
def export_markdown(answer: str, prompt: str, model_id: str, model_name: str) -> BytesIO:
    md = [
        f"# Réponse IA\n",
        f"**Modèle :** `{model_name}` (`{model_id}`)  \n",
        f"**Date :** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n",
        "\n---\n"
    ]
    if prompt:
        md.append("## Prompt\n")
        md.append(f"{prompt}\n\n")
    md.append("## Réponse\n")
    md.append(f"{answer}\n")
    content = "\n".join(md).encode("utf-8")
    return BytesIO(content)

def export_docx(answer: str, prompt: str, model_id: str, model_name: str) -> BytesIO:
    doc = Document()
    doc.add_heading("Réponse IA", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Modèle : {model_name} ({model_id}) – ")
    run.bold = True
    p.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    if prompt:
        doc.add_heading("Prompt", level=2)
        doc.add_paragraph(prompt)

    doc.add_heading("Réponse", level=2)
    for line in answer.splitlines():
        doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_pdf(answer: str, prompt: str, model_id: str, model_name: str) -> BytesIO:
    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 12)
            self.cell(0, 10, "Réponse IA", ln=True, align="C")
            self.ln(2)

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "", 10)
    meta = f"Modèle : {model_name} ({model_id})    Date : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    pdf.multi_cell(0, 6, meta)
    pdf.ln(3)

    if prompt:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Prompt", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, prompt)
        pdf.ln(2)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, "Réponse", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for para in answer.split("\n\n"):
        pdf.multi_cell(0, 6, para.strip())
        pdf.ln(1)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

def save_to_disk(answer: str, prompt: str, model_id: str, model_name: str):
    """💾 Sauvegarde automatique dans /exports sous 3 formats"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"{timestamp}_{model_name.replace(' ', '_')}"

    # Markdown
    md_path = EXPORT_DIR / f"{base_name}.md"
    md_path.write_text(
        f"# Prompt\n{prompt}\n\n# Réponse\n{answer}",
        encoding="utf-8"
    )

    # DOCX
    doc_buf = export_docx(answer, prompt, model_id, model_name)
    (EXPORT_DIR / f"{base_name}.docx").write_bytes(doc_buf.getvalue())

    # PDF
    pdf_buf = export_pdf(answer, prompt, model_id, model_name)
    (EXPORT_DIR / f"{base_name}.pdf").write_bytes(pdf_buf.getvalue())

    st.info(f"🗂️ Sauvegardé dans le dossier **exports/** sous le nom `{base_name}`.")

# -------------------------------
# API OpenRouter – modèles gratuits
# -------------------------------
@st.cache_data(ttl=3600)
def fetch_models():
    url = "https://openrouter.ai/api/v1/models"
    headers = {"Content-Type": "application/json"}
    models = []
    try:
        r = requests.get(url, headers=headers)
        r.raise_for_status()
        data = r.json().get('data', [])
        for m in data:
            pricing = m.get('pricing', {})
            if float(pricing.get('prompt', 0)) == 0 and float(pricing.get('completion', 0)) == 0:
                if any(x in m.get('id', '').lower() or x in m.get('name', '').lower()
                       for x in ['uncensored', 'venice', 'dolphin', 'mai', 'lumimaid']):
                    model_type = "Texte (non censuré)"
                else:
                    model_type = "Texte"
                models.append({
                    'id': m.get('id'),
                    'name': m.get('name'),
                    'type': model_type,
                    'desc': (m.get('description') or '-')[:96]
                })
        return models
    except Exception as e:
        st.error(f"Erreur chargement modèles API OpenRouter : {e}")
        return []

models = fetch_models()

# -------------------------------
# Interface principale
# -------------------------------
st.sidebar.header("🔑 Clé API OpenRouter")
api_key = st.sidebar.text_input("Entre ta clé API OpenRouter :", type="password")
if not api_key:
    st.sidebar.warning("Ajoute ta clé API OpenRouter pour interroger les modèles.")

if models:
    models_names = [f"{m['name']} – {m['type']}" for m in models]
    model_idx = st.selectbox("Choisis ton modèle IA :", models_names)
    model_selected = models[models_names.index(model_idx)]

    st.write(f"**Type :** {model_selected['type']}")
    st.write(f"**Description :** {model_selected['desc']}")

    prompt = st.text_area("Prompt à envoyer :", "")
    if st.button("Envoyer le prompt"):
        if not api_key:
            st.error("Tu dois saisir une clé API OpenRouter dans la barre latérale.")
        elif not prompt.strip():
            st.warning("Entre un prompt puis réessaie !")
        else:
            url = "https://openrouter.ai/api/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost",
                "X-Title": "Streamlit OpenRouter"
            }
            payload = {
                "model": model_selected['id'],
                "messages": [{"role": "user", "content": prompt.strip()}],
            }
            try:
                r = requests.post(url, json=payload, headers=headers, timeout=30)
                r.raise_for_status()
                result = r.json()
                answer = result["choices"][0]["message"]["content"]
                st.success("Réponse de l'IA :")
                st.write(answer)

                # Sauvegarde auto
                save_to_disk(answer, prompt.strip(), model_selected['id'], model_selected['name'])

                # Stocke pour export manuel
                st.session_state["last_answer"] = {
                    "text": answer,
                    "prompt": prompt.strip(),
                    "model_id": model_selected['id'],
                    "model_name": model_selected['name']
                }
            except Exception as e:
                st.error(f"Erreur API : {e}")
else:
    st.warning("Aucun modèle gratuit trouvé.")

# -------------------------------
# ⬇️ Exports manuels
# -------------------------------
if "last_answer" in st.session_state:
    st.markdown("### ⬇️ Exporter la réponse")
    ans = st.session_state["last_answer"]

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button(
            "💾 Exporter en Markdown",
            data=export_markdown(ans["text"], ans["prompt"], ans["model_id"], ans["model_name"]),
            file_name="reponse_ia.md",
            mime="text/markdown"
        )
    with col2:
        st.download_button(
            "💾 Exporter en Word (.docx)",
            data=export_docx(ans["text"], ans["prompt"], ans["model_id"], ans["model_name"]),
            file_name="reponse_ia.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col3:
        st.download_button(
            "💾 Exporter en PDF",
            data=export_pdf(ans["text"], ans["prompt"], ans["model_id"], ans["model_name"]),
            file_name="reponse_ia.pdf",
            mime="application/pdf"
        )

st.caption("Source : openrouter.ai – liste dynamique, prompt direct !")
